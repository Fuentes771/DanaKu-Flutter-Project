#!/usr/bin/env python3
"""
Script untuk convert JOBDESK_TIM.md ke format Word (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style the hyperlink
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def process_markdown_to_docx(md_file, output_file):
    """Convert markdown file to Word document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    in_table = False
    table = None
    table_headers = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines in code blocks
        if in_code_block and not line:
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(51, 51, 51)
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            i += 1
            continue
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            i += 1
            continue
        
        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Tables
        if '|' in line and not in_table:
            # Start of table
            table_headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            
            # Check if next line is separator
            if i + 1 < len(lines) and '|' in lines[i + 1] and '-' in lines[i + 1]:
                in_table = True
                table = doc.add_table(rows=1, cols=len(table_headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                header_cells = table.rows[0].cells
                for idx, header in enumerate(table_headers):
                    header_cells[idx].text = header
                    # Bold headers
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                i += 2  # Skip header and separator line
                continue
        
        if in_table and '|' in line:
            if line.strip() == '':
                in_table = False
                table = None
                table_headers = []
                i += 1
                continue
            
            cells_data = [cell.strip() for cell in line.split('|') if cell.strip()]
            if len(cells_data) == len(table_headers):
                row_cells = table.add_row().cells
                for idx, cell_data in enumerate(cells_data):
                    row_cells[idx].text = cell_data
            i += 1
            continue
        
        if in_table and not '|' in line:
            in_table = False
            table = None
            table_headers = []
        
        # Bullet points
        if line.startswith('- [ ] ') or line.startswith('- [x] ') or line.startswith('- [X] '):
            checked = '[x]' in line.lower()
            text = line[6:]
            p = doc.add_paragraph(text, style='List Bullet')
            run = p.runs[0]
            if checked:
                run.font.color.rgb = RGBColor(0, 128, 0)
            i += 1
            continue
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Bold text with emoji or special formatting
        if line.startswith('**') and line.endswith('**'):
            text = line.strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
        
        # Regular paragraphs
        if line.strip():
            # Handle links [text](url)
            link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
            if re.search(link_pattern, line):
                p = doc.add_paragraph()
                last_end = 0
                for match in re.finditer(link_pattern, line):
                    # Add text before link
                    if match.start() > last_end:
                        p.add_run(line[last_end:match.start()])
                    # Add hyperlink
                    add_hyperlink(p, match.group(1), match.group(2))
                    last_end = match.end()
                # Add remaining text
                if last_end < len(line):
                    p.add_run(line[last_end:])
            else:
                # Process inline formatting
                p = doc.add_paragraph()
                
                # Handle bold **text**
                parts = re.split(r'(\*\*[^\*]+\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    else:
                        p.add_run(part)
        else:
            # Empty line
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✅ Dokumen Word berhasil dibuat: {output_file}")

if __name__ == '__main__':
    input_file = 'JOBDESK_TIM.md'
    output_file = 'JOBDESK_TIM.docx'
    
    try:
        process_markdown_to_docx(input_file, output_file)
        print(f"\n📄 File '{output_file}' siap dibagikan ke tim!")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
